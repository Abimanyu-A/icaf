import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


class DOCXGenerator:

    def __init__(self, output_dir):
        self.output_dir = output_dir

    def _add_page_number(self, document):

        section = document.sections[0]
        footer = section.footer
        paragraph = footer.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = paragraph.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')

        instrText = OxmlElement('w:instrText')
        instrText.text = "PAGE"

        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')

        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)

    def _add_title(self, doc):

        title = doc.add_heading(
            "Telecom Compliance Automation Framework (TCAF)", 0
        )
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def _add_dut_details(self, doc, context):

        doc.add_heading("DUT DETAILS", level=1)

        table = doc.add_table(rows=4, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        rows = [
            ("Device", context.dut_model),
            ("Serial Number", context.dut_serial),
            ("Firmware Version", context.dut_firmware),
            ("DUT IP Address", context.ssh_ip),
        ]

        for i, (key, value) in enumerate(rows):
            table.rows[i].cells[0].text = key
            table.rows[i].cells[1].text = str(value)

    def _add_itsar(self, doc, context):

        doc.add_heading("ITSAR INFORMATION", level=1)

        table = doc.add_table(rows=2, cols=2)

        table.rows[0].cells[0].text = "Section"
        table.rows[0].cells[1].text = context.itsar_section

        table.rows[1].cells[0].text = "Requirement"
        table.rows[1].cells[1].text = context.itsar_requirement

    def _add_requirement(self, doc):

        description = """
The CPE shall communicate with authenticated management entities only.
The protocols used for the CPE management shall support mutual
authentication mechanisms using authentication attributes such as
username/password or equivalent mechanisms.
"""

        doc.add_heading("Requirement Description", level=1)
        doc.add_paragraph(description)

    def _add_test_cases(self, doc, results):

        doc.add_heading("TEST EXECUTION", level=1)

        for idx, tc in enumerate(results, start=1):

            doc.add_heading(
                f"Test Case {idx}: {tc.name}",
                level=2
            )

            doc.add_paragraph(f"Description: {tc.description}")

            p = doc.add_paragraph("Result: ")
            run = p.add_run(tc.status)

            if tc.status.upper() == "PASS":
                run.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
            else:
                run.font.highlight_color = WD_COLOR_INDEX.RED

            doc.add_paragraph("")

            # Evidence images
            for evidence in tc.evidence:

                if os.path.exists(evidence):

                    doc.add_picture(
                        evidence,
                        width=Inches(6)
                    )

                    caption = doc.add_paragraph(
                        f"Evidence Screenshot - {os.path.basename(evidence)}"
                    )
                    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER

            doc.add_paragraph("")

    def _add_result_summary(self, doc, results):

        doc.add_heading("TEST CASE RESULT SUMMARY", level=1)

        table = doc.add_table(
            rows=len(results) + 1,
            cols=4
        )

        headers = [
            "SL No",
            "TEST CASE NAME",
            "PASS/FAIL",
            "Remarks"
        ]

        for i, h in enumerate(headers):
            table.rows[0].cells[i].text = h

        for i, tc in enumerate(results, start=1):

            table.rows[i].cells[0].text = str(i)
            table.rows[i].cells[1].text = tc.name
            table.rows[i].cells[2].text = tc.status
            table.rows[i].cells[3].text = getattr(tc, "remarks", "")

    def generate(self, context, results):

        report_path = os.path.join(
            self.output_dir,
            "tcaf_report.docx"
        )

        doc = Document()

        self._add_page_number(doc)
        self._add_title(doc)
        doc.add_paragraph("")

        self._add_dut_details(doc, context)
        doc.add_paragraph("")

        self._add_itsar(doc, context)
        doc.add_paragraph("")

        self._add_requirement(doc)
        doc.add_paragraph("")

        self._add_test_cases(doc, results)
        doc.add_paragraph("")

        self._add_result_summary(doc, results)

        doc.save(report_path)

        return report_path