"""
clause_1_6_1_report.py

Full ITSAR 1.6.1 Cryptographic Based Secure Communication report.
Covers all 10 test cases:
  SSH  : TC1 Cipher Detection, TC2 Secure Comm, TC3 Weak Cipher, TC4 None-Cipher
  HTTPS: TC5 Cipher Detection, TC6 Secure Comm, TC7 Weak Cipher, TC8 NULL Cipher
  SNMP : TC9 Version Check (v1/v2c disabled), TC10 Secure Comm (authPriv only)
"""

import os
from datetime import datetime

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL

from icaf.reporting.front_page import add_front_page


# ─────────────────────────────────────────────────────────────────────────────
# Low-level formatting helpers
# ─────────────────────────────────────────────────────────────────────────────

def _add_itsar_heading(doc, text: str, level: int = 2):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(16 if level == 1 else 14)
    run.font.color.rgb = RGBColor(75, 0, 130)
    para.paragraph_format.space_before = Pt(14)
    para.paragraph_format.space_after = Pt(2)
    p = para._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), "4B0082")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return para


def _add_itsar_subheading(doc, text: str):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(75, 0, 130)
    para.paragraph_format.space_before = Pt(10)
    para.paragraph_format.space_after = Pt(4)
    return para


def _add_bold_paragraph(doc, text: str):
    p = doc.add_paragraph()
    p.add_run(text).bold = True
    return p


def _keep_with_next(para):
    para.paragraph_format.keep_with_next = True
    para.paragraph_format.keep_together = True
    return para


def _add_grey_box(doc, text: str):
    """Grey-background bordered box for commands / terminal output."""
    table = doc.add_table(rows=1, cols=1)
    table.allow_autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:type"), "dxa")
    tblW.set(qn("w:w"), str(int(5.8 * 1440)))
    tblPr.append(tblW)
    cell = table.cell(0, 0)
    cell.text = str(text) if text else ""
    shading = parse_xml(r'<w:shd {} w:fill="F2F2F2"/>'.format(nsdecls("w")))
    cell._element.get_or_add_tcPr().append(shading)
    borders = parse_xml(r'''
        <w:tcBorders {}>
            <w:top    w:val="single" w:sz="4" w:color="BFBFBF"/>
            <w:left   w:val="single" w:sz="4" w:color="BFBFBF"/>
            <w:bottom w:val="single" w:sz="4" w:color="BFBFBF"/>
            <w:right  w:val="single" w:sz="4" w:color="BFBFBF"/>
        </w:tcBorders>
    '''.format(nsdecls("w")))
    cell._element.get_or_add_tcPr().append(borders)
    return table


def _add_screenshot_block(doc, title: str, image_path: str):
    """Purple-bordered evidence block with title + image."""
    TABLE_WIDTH = Inches(7.8)
    IMAGE_WIDTH = Inches(6.2)
    table = doc.add_table(rows=2, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.allow_autofit = False
    table.columns[0].width = TABLE_WIDTH
    for row in table.rows:
        cell = row.cells[0]
        cell.width = TABLE_WIDTH
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "F3ECFA")
        tcPr.append(shd)
        cell.top_margin = Inches(0.15)
        cell.bottom_margin = Inches(0.15)
        cell.left_margin = Inches(0.25)
        cell.right_margin = Inches(0.25)
    p_title = table.cell(0, 0).paragraphs[0]
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_title.add_run(title)
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(75, 0, 130)
    p_img = table.cell(1, 0).paragraphs[0]
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_img.add_run().add_picture(image_path, width=IMAGE_WIDTH)
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "12")
        border.set(qn("w:color"), "4B0082")
        tblBorders.append(border)
    tblPr.append(tblBorders)
    return table


def _add_two_col_cipher_table(doc, strong_items, weak_items,
                               strong_header="Strong", weak_header="Weak"):
    strong_items = strong_items or ["None"]
    weak_items   = weak_items   or ["None"]
    t = doc.add_table(rows=2, cols=2)
    t.style = "Table Grid"
    t.cell(0, 0).text = strong_header
    t.cell(0, 1).text = weak_header
    t.cell(1, 0).text = "\n".join(strong_items)
    t.cell(1, 1).text = "\n".join(weak_items)
    return t


def _prevent_row_split(table):
    for row in table.rows:
        trPr = row._tr.get_or_add_trPr()
        cs = OxmlElement("w:cantSplit")
        trPr.append(cs)


def _style_header_row(table, headers, bg="4B0082"):
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), bg)
        tcPr.append(shd)


def _pad_data_rows(table):
    for row in table.rows[1:]:
        for cell in row.cells:
            cell.top_margin    = Inches(0.1)
            cell.bottom_margin = Inches(0.1)
            cell.left_margin   = Inches(0.1)
            cell.right_margin  = Inches(0.1)


def _spacer(doc, pts=12):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(pts)


# ─────────────────────────────────────────────────────────────────────────────
# Main report class
# ─────────────────────────────────────────────────────────────────────────────

class Clause161Report:
    """
    Generates the full ITSAR 1.6.1 Word report covering all 10 test cases.

    context.scan_results must contain keys set by clause_1_6_1/clause.py:
        nmap, ssh_applicable, https_applicable, snmp_applicable,
        cipher, ssh, weak_cipher, none_cipher,
        https_cipher, https, https_weak_cipher, https_null,
        snmp_v1v2, snmp_v3, oem, dut_info
    """

    HEADERS = ["SL. No", "TEST CASE NAME", "PASS / FAIL", "Remarks"]

    def __init__(self, context, results):
        self.context    = context
        self.results    = results
        self.output_dir = context.evidence.run_dir

    # ── Public entry point ───────────────────────────────────────────────────

    def generate(self) -> str:
        r   = self.context.scan_results
        doc = Document()

        dut    = r["dut_info"]
        ssh_ok = r.get("ssh_applicable",   False)
        tls_ok = r.get("https_applicable", False)
        snmp_ok = r.get("snmp_applicable", False)

        # Compute all TC results
        tc1  = r["cipher"].get("result", "NA")             if ssh_ok  else "NA"
        tc2  = r["ssh"].get("final_result", "NA")          if ssh_ok  else "NA"
        tc3  = self._weak_result(r["weak_cipher"])          if ssh_ok  else "NA"
        tc4  = r["none_cipher"].get("result", "NA")        if ssh_ok  else "NA"
        tc5  = r["https_cipher"].get("result", "NA")       if tls_ok  else "NA"
        tc6  = r["https"].get("final_result", "NA")        if tls_ok  else "NA"
        tc7  = self._https_weak_result(r["https_weak_cipher"]) if tls_ok else "NA"
        tc8  = r["https_null"].get("final_result", "NA")   if tls_ok  else "NA"
        tc9  = r["snmp_v1v2"].get("final_result", "NA")    if snmp_ok else "NA"
        tc10 = r["snmp_v3"].get("final_result",  "NA")     if snmp_ok else "NA"

        all_tcs = [t for t in [tc1,tc2,tc3,tc4,tc5,tc6,tc7,tc8,tc9,tc10] if t != "NA"]
        final   = "PASS" if all_tcs and all(t == "PASS" for t in all_tcs) else "FAIL"

        # Front page
        add_front_page(doc, meta={
            "dut_name":      dut["dut_name"],
            "dut_version":   dut["dut_version"],
            "os_hash":       dut["os_hash"],
            "config_hash":   dut["config_hash"],
            "start_time":    str(self.context.start_time),
            "end_time":      "",
            "final_result":  final,
            "itsar_id":      "ITSAR 2.6.1",
            "itsar_version": "2.6.1",
        })

        sec = [0]  # mutable counter

        def next_sec():
            sec[0] += 1
            return sec[0]

        # ── Intro sections ───────────────────────────────────────────────────
        self._intro(doc, next_sec, r)

        # ── SSH sections ─────────────────────────────────────────────────────
        if ssh_ok:
            self._ssh_execution(doc, next_sec, r, tc1, tc2, tc3, tc4)
        else:
            _add_itsar_heading(doc, f"{next_sec()}. Test Execution For SSH")
            doc.add_paragraph(
                "During the service discovery phase using Nmap scan, the SSH service "
                "(port 22) was not detected on the DUT. "
                "Therefore, SSH-based secure communication validation tests were not executed."
            )

        # ── HTTPS sections ───────────────────────────────────────────────────
        if tls_ok:
            self._https_execution(doc, next_sec, r, tc5, tc6, tc7, tc8)
        else:
            _add_itsar_heading(doc, f"{next_sec()}. Test Execution For HTTPS")
            doc.add_paragraph(
                "During the service discovery phase using Nmap scan, the HTTPS service "
                "(port 443) was not detected on the DUT. "
                "Therefore, HTTPS-based secure communication validation tests were not executed."
            )

        # ── SNMP sections ────────────────────────────────────────────────────
        if snmp_ok:
            self._snmp_execution(doc, next_sec, r, tc9, tc10)
        else:
            _add_itsar_heading(doc, f"{next_sec()}. Test Execution For SNMP")
            doc.add_paragraph(
                "During the service discovery phase using Nmap scan, the SNMP service "
                "(port 161) was not detected on the DUT. "
                "Therefore, SNMP-based secure communication validation tests were not executed."
            )

        # ── Final result table ───────────────────────────────────────────────
        self._result_table(doc, next_sec, tc1,tc2,tc3,tc4, tc5,tc6,tc7,tc8, tc9,tc10)

        # ── Save ─────────────────────────────────────────────────────────────
        fname = (
            "ITSAR_2.6.1_Cryptographic_Secure_Communication_Report_"
            f"{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        )
        path = os.path.join(self.output_dir, fname)
        doc.save(path)
        return path

    # ── Intro (nmap + preconditions) ─────────────────────────────────────────

    def _intro(self, doc, next_sec, r):
        nmap = r["nmap"]

        _add_itsar_heading(doc, f"{next_sec()}. Access Authorization")
        doc.add_paragraph(
            "Access authorization ensures that only authenticated and authorized entities "
            "can access network resources and management interfaces on the DUT."
        )

        _add_itsar_heading(doc, f"{next_sec()}. Cryptographic Based Secure Communication")
        doc.add_paragraph(
            "This section verifies that all management and data-plane communications on the DUT "
            "are protected using strong, approved cryptographic algorithms and protocols."
        )

        _add_itsar_heading(doc, f"{next_sec()}. Requirement Description")
        doc.add_paragraph(
            "The DUT shall enforce cryptographic controls for all supported management protocols "
            "(SSH, HTTPS, SNMP). Weak algorithms, insecure protocol versions, and NULL/none cipher "
            "suites must be rejected."
        )

        s = next_sec()
        _add_itsar_heading(doc, f"{s}. DUT Configuration")
        _add_itsar_subheading(doc, f"{s}.1 Service Discovery (Nmap Scan)")
        doc.add_paragraph(
            "An Nmap scan was performed to discover active services on the DUT before running "
            "protocol-specific test cases."
        )
        _add_bold_paragraph(doc, "Execution Command (TCP):")
        _add_grey_box(doc, nmap.get("user_input_tcp_ports", nmap.get("user_input", "N/A")))
        doc.add_paragraph("")
        _add_bold_paragraph(doc, "Executed Command Output (TCP):")
        _add_grey_box(doc, nmap.get("terminal_output_tcp_ports", nmap.get("terminal_output", "No output")))
        doc.add_paragraph("")

        if nmap.get("user_input_udp_ports"):
            _add_bold_paragraph(doc, "Execution Command (UDP):")
            _add_grey_box(doc, nmap["user_input_udp_ports"])
            doc.add_paragraph("")
            _add_bold_paragraph(doc, "Executed Command Output (UDP):")
            _add_grey_box(doc, nmap.get("terminal_output_udp_ports", "No output"))
            doc.add_paragraph("")

        if nmap.get("screenshot") and os.path.exists(nmap["screenshot"]):
            _add_screenshot_block(doc, "DUT Configuration: Nmap Scan", nmap["screenshot"])

        _add_itsar_subheading(doc, f"{s}.2 Detected Services")
        services = []
        for svc in ("SSH", "HTTPS", "SNMP"):
            status = "Detected" if nmap.get(svc) else "Not Detected"
            services.append(f"• {svc}: {status}")
        doc.add_paragraph("\n".join(services))

        _add_itsar_heading(doc, f"{next_sec()}. Pre-conditions")
        for bullet in [
            "DUT is powered on and accessible over the network.",
            "Tester system has SSH, OpenSSL, Nmap, and snmpwalk tools installed.",
            "Valid credentials for SSH, HTTPS, and SNMP are available.",
            "Network capture interface is configured (for SNMP Wireshark evidence).",
            "OAM/OEM declaration document is available if applicable.",
        ]:
            doc.add_paragraph(f"• {bullet}")

        _add_itsar_heading(doc, f"{next_sec()}. Test Objective")
        doc.add_paragraph(
            "Verify that all management protocols on the DUT enforce strong cryptographic controls, "
            "reject weak or deprecated algorithms, and do not permit unencrypted communication."
        )

        _add_itsar_heading(doc, f"{next_sec()}. Expected Results")
        doc.add_paragraph(
            "The DUT shall: support only strong cryptographic algorithms; reject weak cipher "
            "negotiation attempts; refuse SSH and HTTPS connections with none/NULL ciphers; "
            "disable SNMPv1/v2c; and enforce SNMPv3 authPriv only."
        )

    # ── SSH execution (TC1–TC4) ───────────────────────────────────────────────

    def _ssh_execution(self, doc, next_sec, r, tc1, tc2, tc3, tc4):
        cipher   = r["cipher"]
        ssh      = r["ssh"]
        weak     = r["weak_cipher"]
        none_c   = r["none_cipher"]

        s = next_sec()
        _add_itsar_heading(doc, f"{s}. Test Execution For SSH")

        # ── TC1: Cipher Detection ────────────────────────────────────────────
        _add_itsar_subheading(doc, f"{s}.1  TC1 — SSH Secure Cryptographic Algorithms")
        _add_bold_paragraph(doc, "a) Test Case Name:")
        doc.add_paragraph("TC1_SSH_SECURE_CRYPTOGRAPHIC_ALGORITHMS")
        _add_bold_paragraph(doc, "b) Test Case Description:")
        doc.add_paragraph(
            "Enumerate all SSH encryption, MAC, key-exchange, and host-key algorithms "
            "advertised by the DUT and classify each as strong or weak."
        )
        _add_bold_paragraph(doc, "c) Execution Steps:")
        for step in [
            "• Run nmap --script ssh2-enum-algos against the DUT.",
            "• Capture the advertised algorithm lists.",
            "• Classify each algorithm against the approved ITSAR list.",
        ]:
            doc.add_paragraph(step)
        _add_bold_paragraph(doc, "Execution Command:")
        _add_grey_box(doc, cipher.get("user_input", "N/A"))
        doc.add_paragraph("")
        _add_bold_paragraph(doc, "Executed Command Output:")
        _add_grey_box(doc, cipher.get("terminal_output", "No output"))
        doc.add_paragraph("")

        details = cipher.get("details", {})
        for label, key in [
            ("Encryption Algorithms", "encryption"),
            ("MAC Algorithms", "mac"),
            ("Key Exchange Algorithms", "kex"),
            ("Host Key Algorithms", "host_key"),
        ]:
            h = _add_itsar_subheading(doc, label)
            _keep_with_next(h)
            _add_two_col_cipher_table(
                doc,
                details.get(key, {}).get("strong", []),
                details.get(key, {}).get("weak",   []),
                "Strong", "Weak"
            )
            doc.add_paragraph("")

        if cipher.get("screenshot") and os.path.exists(cipher["screenshot"]):
            _add_screenshot_block(doc, "TC1 — SSH Algorithm Enumeration", cipher["screenshot"])
        _add_bold_paragraph(doc, "d) Test Observations:")
        doc.add_paragraph(
            "All advertised SSH algorithms have been enumerated. "
            + ("No weak algorithms detected." if tc1 == "PASS"
               else "Weak algorithm(s) detected — see table above.")
        )
        _spacer(doc)

        # ── TC2: SSH Secure Communication ────────────────────────────────────
        _add_itsar_subheading(doc, f"{s}.2  TC2 — SSH Secure Communication")
        _add_bold_paragraph(doc, "a) Test Case Name:")
        doc.add_paragraph("TC2_SSH_SECURE_COMMUNICATION")
        _add_bold_paragraph(doc, "b) Test Case Description:")
        doc.add_paragraph(
            "Verify that an SSH session to the DUT uses a strong key-exchange algorithm, "
            "encryption cipher, and MAC. Wireshark/packet evidence is captured."
        )
        _add_bold_paragraph(doc, "c) Execution Steps:")
        for step in [
            "• Connect to the DUT via SSH.",
            "• Capture the negotiated KEX, cipher, and MAC from the session output.",
            "• Validate each negotiated algorithm against the NIST-approved list.",
        ]:
            doc.add_paragraph(step)
        _add_bold_paragraph(doc, "Execution Command:")
        _add_grey_box(doc, ssh.get("user_input", "N/A"))
        doc.add_paragraph("")

        crypto = ssh.get("crypto_details", {})
        tbl = doc.add_table(rows=4, cols=2)
        tbl.style = "Table Grid"
        tbl.cell(0,0).text = "Protocol";               tbl.cell(0,1).text = crypto.get("protocol", "Not Found")
        tbl.cell(1,0).text = "Encryption Algorithm";   tbl.cell(1,1).text = crypto.get("cipher",   "Not Found")
        tbl.cell(2,0).text = "Key Exchange Algorithm"; tbl.cell(2,1).text = crypto.get("kex",       "Not Found")
        tbl.cell(3,0).text = "Host Key Algorithm";     tbl.cell(3,1).text = crypto.get("host_key",  "Not Found")
        doc.add_paragraph("")

        if ssh.get("screenshots"):
            for idx, img in enumerate(ssh["screenshots"]):
                if os.path.exists(img):
                    _add_screenshot_block(doc, f"TC2 — SSH Secure Communication Evidence {idx+1}", img)
        _add_bold_paragraph(doc, "d) Test Observations:")
        doc.add_paragraph(
            "SSH session established successfully with strong cryptographic parameters."
            if tc2 == "PASS" else
            "SSH session uses weak or non-compliant cryptographic parameters."
        )
        _spacer(doc)

        # ── TC3: Weak Cipher Negotiation ─────────────────────────────────────
        _add_itsar_subheading(doc, f"{s}.3  TC3 — SSH Weak Cipher Negotiation")
        _add_bold_paragraph(doc, "a) Test Case Name:")
        doc.add_paragraph("TC3_SSH_WEAK_CIPHER_REJECTION")
        _add_bold_paragraph(doc, "b) Test Case Description:")
        doc.add_paragraph(
            "Attempt to negotiate each detected weak SSH algorithm individually. "
            "The DUT must reject all such attempts."
        )
        _add_bold_paragraph(doc, "c) Execution Steps:")
        for step in [
            "• For each weak algorithm found in TC1, attempt SSH connection forcing that algorithm.",
            "• Record whether the connection was accepted or rejected.",
            "• A PASS requires all weak algorithms to be rejected.",
        ]:
            doc.add_paragraph(step)

        for res in weak.get("results", []):
            algo      = res.get("algorithm", "Unknown")
            algo_type = res.get("type",      "Unknown")
            negotiated = res.get("negotiated", False)
            cmd       = res.get("command",   "N/A")
            out       = res.get("terminal_output", "No output")
            _add_bold_paragraph(doc, f"Weak {algo_type.upper()} Attempt: {algo}")
            _add_grey_box(doc, cmd)
            doc.add_paragraph("")
            _add_grey_box(doc, out or "No output")
            doc.add_paragraph("")
            if res.get("screenshot") and os.path.exists(res["screenshot"]):
                title = f"TC3 — SSH Weak {algo_type.upper()} Attempt: {algo}"
                _add_screenshot_block(doc, title, res["screenshot"])
                _add_bold_paragraph(doc, "Explanation:")
                neg_text = "successfully negotiated" if negotiated else "rejected by the DUT"
                doc.add_paragraph(
                    f"The weak {algo_type} algorithm {algo} was {neg_text}. "
                    + ("This is a security violation." if negotiated else
                       "The DUT correctly rejects this insecure algorithm.")
                )
            _spacer(doc, 8)

        _add_bold_paragraph(doc, "d) Test Observations:")
        doc.add_paragraph(
            "All weak SSH algorithm negotiation attempts were rejected by the DUT."
            if tc3 == "PASS" else
            "One or more weak SSH algorithms were successfully negotiated — security violation."
        )
        _spacer(doc)

        # ── TC4: None-Cipher Rejection ───────────────────────────────────────
        _add_itsar_subheading(doc, f"{s}.4  TC4 — SSH None-Cipher Rejection")
        _add_bold_paragraph(doc, "a) Test Case Name:")
        doc.add_paragraph("TC4_SSH_NO_ENCRYPTION_REJECTION")
        _add_bold_paragraph(doc, "b) Test Case Description:")
        doc.add_paragraph(
            "Verify that the DUT refuses SSH connections that explicitly request "
            "Ciphers=none (unencrypted communication)."
        )
        _add_bold_paragraph(doc, "c) Execution Steps:")
        for step in [
            "• Attempt SSH connection with -o Ciphers=none.",
            "• Observe client/server output for rejection message.",
            "• Verify 'none' cipher does not appear in the DUT's advertised algorithm list.",
        ]:
            doc.add_paragraph(step)
        _add_bold_paragraph(doc, "Execution Command:")
        _add_grey_box(doc, none_c.get("user_input", "N/A"))
        doc.add_paragraph("")
        _add_bold_paragraph(doc, "Executed Command Output:")
        _add_grey_box(doc, none_c.get("terminal_output", "No output"))
        doc.add_paragraph("")

        if none_c.get("screenshot") and os.path.exists(none_c["screenshot"]):
            remarks      = none_c.get("remarks", "")
            none_present = none_c.get("None_cipher_exist", False)
            overview = (
                f"The SSH connection attempt using Ciphers=none resulted in: {remarks}. "
            )
            if none_present:
                overview += (
                    "The DUT advertises 'none' cipher in its encryption list — "
                    "critical security vulnerability."
                )
            else:
                overview += "The DUT does not advertise the 'none' cipher — expected secure behaviour."
            _add_screenshot_block(doc, "TC4 — SSH None-Cipher Attempt", none_c["screenshot"])
            _add_bold_paragraph(doc, "Explanation of Above Screenshot:")
            doc.add_paragraph(overview)

        _add_bold_paragraph(doc, "d) Test Observations:")
        result  = none_c.get("result", "UNKNOWN")
        remarks = none_c.get("remarks", "")
        if result == "PASS":
            doc.add_paragraph(
                f"The SSH none-cipher attempt was rejected. {remarks}. "
                "The DUT complies with ITSAR cryptographic requirements — marked PASS."
            )
        else:
            doc.add_paragraph(
                f"The SSH none-cipher attempt indicates: {remarks}. "
                "Unencrypted communication may be possible — marked FAIL."
            )
        _add_bold_paragraph(doc, "e) Evidence Provided:")
        doc.add_paragraph("Screenshots and command outputs provided.")
        _spacer(doc)

    # ── HTTPS execution (TC5–TC8) ─────────────────────────────────────────────

    def _https_execution(self, doc, next_sec, r, tc5, tc6, tc7, tc8):
        hcipher = r["https_cipher"]
        https   = r["https"]
        hweak   = r["https_weak_cipher"]
        hnull   = r["https_null"]

        s = next_sec()
        _add_itsar_heading(doc, f"{s}. Test Execution For HTTPS")

        # ── TC5: TLS Cipher Detection ─────────────────────────────────────────
        _add_itsar_subheading(doc, f"{s}.1  TC5 — HTTPS TLS Cipher Detection")
        _add_bold_paragraph(doc, "a) Test Case Name:")
        doc.add_paragraph("TC5_HTTPS_SECURE_CRYPTOGRAPHIC_ALGORITHMS")
        _add_bold_paragraph(doc, "b) Test Case Description:")
        doc.add_paragraph(
            "Enumerate all TLS cipher suites supported by the DUT HTTPS service "
            "for TLS 1.2 and TLS 1.3, classifying each as strong or weak."
        )
        _add_bold_paragraph(doc, "c) Execution Steps:")
        for step in [
            "• Run nmap --script ssl-enum-ciphers against port 443.",
            "• Enumerate TLS 1.2 and TLS 1.3 cipher suites.",
            "• Classify each cipher against the ITSAR approved list.",
        ]:
            doc.add_paragraph(step)
        _add_bold_paragraph(doc, "Execution Command:")
        _add_grey_box(doc, hcipher.get("user_input", "N/A"))
        doc.add_paragraph("")
        _add_bold_paragraph(doc, "Executed Command Output:")
        _add_grey_box(doc, hcipher.get("terminal_output", "No output"))
        doc.add_paragraph("")

        details = hcipher.get("details", {})
        for ver, ver_key in [("TLSv1.2", "TLSv1.2"), ("TLSv1.3", "TLSv1.3")]:
            ver_data = details.get(ver_key, {})
            for label, key in [
                (f"{ver} — Cipher Suites", "ciphers"),
                (f"{ver} — Encryption",    "encryption"),
                (f"{ver} — MAC",           "mac"),
                (f"{ver} — KEX",           "kex"),
            ]:
                block = ver_data.get(key, {})
                if isinstance(block, dict) and ("strong" in block or "weak" in block):
                    h = _add_itsar_subheading(doc, label)
                    _keep_with_next(h)
                    _add_two_col_cipher_table(
                        doc,
                        block.get("strong", []),
                        block.get("weak",   []),
                        "Strong", "Weak"
                    )
                    doc.add_paragraph("")

        if hcipher.get("screenshot") and os.path.exists(hcipher["screenshot"]):
            _add_screenshot_block(doc, "TC5 — HTTPS TLS Cipher Enumeration", hcipher["screenshot"])
        _add_bold_paragraph(doc, "d) Test Observations:")
        doc.add_paragraph(
            "All supported TLS cipher suites enumerated. "
            + ("No weak ciphers detected." if tc5 == "PASS" else "Weak cipher(s) detected.")
        )
        _spacer(doc)

        # ── TC6: HTTPS Secure Communication ──────────────────────────────────
        _add_itsar_subheading(doc, f"{s}.2  TC6 — HTTPS Secure Communication")
        _add_bold_paragraph(doc, "a) Test Case Name:")
        doc.add_paragraph("TC6_HTTPS_SECURE_COMMUNICATION")
        _add_bold_paragraph(doc, "b) Test Case Description:")
        doc.add_paragraph(
            "Verify that an HTTPS connection to the DUT negotiates a strong TLS protocol "
            "version and cipher suite."
        )
        _add_bold_paragraph(doc, "c) Execution Steps:")
        for step in [
            "• Connect using openssl s_client -connect <DUT>:443.",
            "• Capture the negotiated protocol version and cipher.",
            "• Validate against the ITSAR-approved list.",
        ]:
            doc.add_paragraph(step)
        _add_bold_paragraph(doc, "Execution Command:")
        _add_grey_box(doc, https.get("user_input", "N/A"))
        doc.add_paragraph("")
        _add_bold_paragraph(doc, "Executed Command Output:")
        _add_grey_box(doc, https.get("terminal_output", "No output"))
        doc.add_paragraph("")

        crypto = https.get("crypto_details", {})
        tbl = doc.add_table(rows=2, cols=2)
        tbl.style = "Table Grid"
        tbl.cell(0,0).text = "Protocol";             tbl.cell(0,1).text = crypto.get("protocol", "Not Found")
        tbl.cell(1,0).text = "Encryption Algorithm"; tbl.cell(1,1).text = crypto.get("cipher",   "Not Found")
        doc.add_paragraph("")

        if https.get("screenshots"):
            for idx, img in enumerate(https["screenshots"]):
                if os.path.exists(img):
                    _add_screenshot_block(doc, f"TC6 — HTTPS Secure Communication Evidence {idx+1}", img)
        _add_bold_paragraph(doc, "d) Test Observations:")
        doc.add_paragraph(
            "HTTPS session established with strong TLS configuration."
            if tc6 == "PASS" else
            "HTTPS session uses weak or non-compliant TLS configuration."
        )
        _spacer(doc)

        # ── TC7: HTTPS Weak Cipher Negotiation ───────────────────────────────
        _add_itsar_subheading(doc, f"{s}.3  TC7 — HTTPS Weak Cipher Negotiation")
        _add_bold_paragraph(doc, "a) Test Case Name:")
        doc.add_paragraph("TC7_HTTPS_WEAK_CIPHER_REJECTION")
        _add_bold_paragraph(doc, "b) Test Case Description:")
        doc.add_paragraph(
            "Attempt to negotiate each detected weak TLS cipher suite against the DUT. "
            "The DUT must reject all such attempts."
        )
        _add_bold_paragraph(doc, "c) Execution Steps:")
        for step in [
            "• For each weak cipher found in TC5, force it via openssl s_client -cipher / -ciphersuites.",
            "• Analyze the handshake output for acceptance or rejection.",
            "• PASS requires all weak ciphers to be rejected.",
        ]:
            doc.add_paragraph(step)

        for res in hweak.get("results", []):
            cipher_name = res.get("cipher",          "Unknown")
            tls_ver     = res.get("tls_version",     "Unknown")
            cmd         = res.get("command",         "N/A")
            out         = res.get("terminal_output", "No output")
            negotiated  = res.get("negotiated",       False)
            neg_text    = "successfully negotiated" if negotiated else "rejected by the DUT"

            _add_bold_paragraph(doc, f"Weak Cipher Attempt ({tls_ver}): {cipher_name}")
            _add_grey_box(doc, cmd)
            doc.add_paragraph("")
            _add_grey_box(doc, out or "No relevant output (cipher likely rejected)")
            doc.add_paragraph("")
            if res.get("screenshot") and os.path.exists(res["screenshot"]):
                _add_screenshot_block(
                    doc,
                    f"TC7 — HTTPS Weak Cipher Attempt ({tls_ver}): {cipher_name}",
                    res["screenshot"]
                )
                _add_bold_paragraph(doc, "Explanation:")
                doc.add_paragraph(
                    f"The weak TLS cipher {cipher_name} ({tls_ver}) was {neg_text}. "
                    + ("Security violation detected." if negotiated else
                       "DUT correctly rejects this cipher.")
                )
            _spacer(doc, 8)

        _add_bold_paragraph(doc, "d) Test Observations:")
        neg_list = [r["cipher"] for r in hweak.get("results", []) if r.get("negotiated")]
        if neg_list:
            doc.add_paragraph(
                f"Weak TLS cipher(s) {', '.join(neg_list)} were successfully negotiated — "
                "security violation. Marked FAIL."
            )
        else:
            doc.add_paragraph(
                "All weak TLS cipher negotiation attempts were rejected by the DUT. Marked PASS."
            )
        _add_bold_paragraph(doc, "e) Evidence Provided:")
        doc.add_paragraph("Screenshots and command outputs provided.")
        _spacer(doc)

        # ── TC8: HTTPS NULL Cipher Rejection ──────────────────────────────────
        _add_itsar_subheading(doc, f"{s}.4  TC8 — HTTPS NULL Cipher Rejection")
        _add_bold_paragraph(doc, "a) Test Case Name:")
        doc.add_paragraph("TC8_HTTPS_NO_ENCRYPTION_REJECTION")
        _add_bold_paragraph(doc, "b) Test Case Description:")
        doc.add_paragraph(
            "Verify that the DUT rejects HTTPS connections using NULL (no-encryption) "
            "cipher suites under TLS 1.2, and that TLS 1.3 enforces strong ciphers."
        )
        _add_bold_paragraph(doc, "c) Execution Steps:")
        for step in [
            "• Attempt HTTPS connection with -cipher NULL -tls1_2 via openssl s_client.",
            "• Verify handshake fails or no cipher is negotiated.",
            "• Connect normally with TLS 1.3 and verify a strong cipher is negotiated.",
        ]:
            doc.add_paragraph(step)

        for ver_key, label in [("tls1_2", "TLS 1.2 NULL Cipher"), ("tls1_3", "TLS 1.3 Strong Cipher")]:
            v = hnull.get(ver_key, {})
            _add_bold_paragraph(doc, f"Execution Command ({label}):")
            _add_grey_box(doc, v.get("command", "N/A"))
            doc.add_paragraph("")
            _add_bold_paragraph(doc, f"Executed Command Output ({label}):")
            _add_grey_box(doc, v.get("output", "No output"))
            doc.add_paragraph("")

            if v.get("screenshot") and os.path.exists(v["screenshot"]):
                remarks = v.get("remarks", "")
                overview = (
                    f"The HTTPS connection attempt using {label} resulted in: {remarks}. "
                    + ("The DUT correctly rejects unencrypted communication."
                       if v.get("result") == "PASS" else
                       "Unencrypted communication may be permitted — security violation.")
                )
                _add_screenshot_block(doc, f"TC8 — HTTPS {label}", v["screenshot"])
                _add_bold_paragraph(doc, "Explanation of Above Screenshot:")
                doc.add_paragraph(overview)
            _spacer(doc, 8)

        _add_bold_paragraph(doc, "d) Test Observations:")
        final_null = hnull.get("final_result", "UNKNOWN")
        if final_null == "PASS":
            doc.add_paragraph(
                "NULL cipher was rejected under TLS 1.2 and TLS 1.3 enforced strong encryption. "
                "DUT complies with ITSAR requirements — marked PASS."
            )
        else:
            doc.add_paragraph(
                "NULL or unencrypted TLS communication may be possible on this DUT. Marked FAIL."
            )
        _add_bold_paragraph(doc, "e) Evidence Provided:")
        doc.add_paragraph("Screenshots and terminal outputs provided.")
        _spacer(doc)

    # ── SNMP execution (TC9–TC10) ─────────────────────────────────────────────

    def _snmp_execution(self, doc, next_sec, r, tc9, tc10):
        v1v2 = r["snmp_v1v2"]
        v3   = r["snmp_v3"]

        s = next_sec()
        _add_itsar_heading(doc, f"{s}. Test Execution For SNMP")

        # ── TC9: SNMP Version Check ───────────────────────────────────────────
        _add_itsar_subheading(doc, f"{s}.1  TC9 — SNMP Version Check (v1/v2c Disabled)")
        _add_bold_paragraph(doc, "a) Test Case Name:")
        doc.add_paragraph("TC9_SNMP_VERSION_CHECK")
        _add_bold_paragraph(doc, "b) Test Case Description:")
        doc.add_paragraph(
            "Verify that the DUT does not respond to SNMPv1 or SNMPv2c requests. "
            "Only SNMPv3 should be active. Any successful v1/v2c response indicates "
            "insecure configuration."
        )
        _add_bold_paragraph(doc, "c) Execution Steps:")
        for step in [
            "• Attempt SNMPv1 walk using default community string (public).",
            "• Attempt SNMPv2c walk using default community string (public).",
            "• Record whether the DUT responds with valid SNMP data.",
            "• A successful response = FAIL (insecure protocol enabled).",
        ]:
            doc.add_paragraph(step)

        # SNMPv1 block
        _add_bold_paragraph(doc, "Execution Command (SNMPv1):")
        _add_grey_box(doc, v1v2.get("user_input_v1", "N/A"))
        doc.add_paragraph("")
        _add_bold_paragraph(doc, "Executed Command Output (SNMPv1):")
        _add_grey_box(doc, v1v2.get("terminal_output_v1", "No output"))
        doc.add_paragraph("")

        v1_ok = v1v2.get("validation_details", {}).get("v1_success", False)
        if v1v2.get("v1_screenshot") and os.path.exists(v1v2["v1_screenshot"]):
            overview = (
                "The DUT responded to SNMPv1 — insecure protocol is enabled (FAIL)."
                if v1_ok else
                "No response to SNMPv1 — insecure protocol correctly disabled (PASS)."
            )
            _add_screenshot_block(doc, "TC9 — SNMPv1 Communication Attempt", v1v2["v1_screenshot"])
            _add_bold_paragraph(doc, "Explanation of Above Screenshot:")
            doc.add_paragraph(overview)

        # SNMPv2c block
        _add_bold_paragraph(doc, "Execution Command (SNMPv2c):")
        _add_grey_box(doc, v1v2.get("user_input_v2c", "N/A"))
        doc.add_paragraph("")
        _add_bold_paragraph(doc, "Executed Command Output (SNMPv2c):")
        _add_grey_box(doc, v1v2.get("terminal_output_v2c", "No output"))
        doc.add_paragraph("")

        v2_ok = v1v2.get("validation_details", {}).get("v2c_success", False)
        if v1v2.get("v2c_screenshot") and os.path.exists(v1v2["v2c_screenshot"]):
            overview = (
                "The DUT responded to SNMPv2c — insecure protocol is enabled (FAIL)."
                if v2_ok else
                "No response to SNMPv2c — insecure protocol correctly disabled (PASS)."
            )
            _add_screenshot_block(doc, "TC9 — SNMPv2c Communication Attempt", v1v2["v2c_screenshot"])
            _add_bold_paragraph(doc, "Explanation of Above Screenshot:")
            doc.add_paragraph(overview)

        _add_bold_paragraph(doc, "d) Test Observations:")
        if tc9 == "PASS":
            doc.add_paragraph(
                "SNMPv1 and SNMPv2c requests received no response. "
                "Insecure SNMP versions are disabled — DUT marked PASS."
            )
        else:
            doc.add_paragraph(
                "SNMPv1 and/or SNMPv2c requests received a valid response. "
                "Insecure SNMP versions are enabled — DUT marked FAIL."
            )
        _add_bold_paragraph(doc, "e) Evidence Provided:")
        doc.add_paragraph("Screenshots of SNMPv1 and SNMPv2c communication attempts provided.")
        _spacer(doc)

        # ── TC10: SNMP Secure Communication ──────────────────────────────────
        _add_itsar_subheading(doc, f"{s}.2  TC10 — SNMP Secure Communication (authPriv Only)")
        _add_bold_paragraph(doc, "a) Test Case Name:")
        doc.add_paragraph("TC10_SNMP_SECURE_COMMUNICATION")
        _add_bold_paragraph(doc, "b) Test Case Description:")
        doc.add_paragraph(
            "Verify that SNMPv3 enforces authPriv (SHA + AES) and rejects weaker security levels "
            "(authNoPriv and noAuthNoPriv)."
        )
        _add_bold_paragraph(doc, "c) Execution Steps:")
        for step in [
            "• Initiate SNMPv3 walk using authPriv (SHA + AES) — should succeed.",
            "• Initiate SNMPv3 walk using authNoPriv — should be rejected.",
            "• Initiate SNMPv3 walk using noAuthNoPriv — should be rejected.",
            "• Capture Wireshark evidence confirming msgFlags.",
        ]:
            doc.add_paragraph(step)

        for mode_key, label, expect_success in [
            ("authPriv",     "authPriv (SHA + AES)",         True),
            ("authNoPriv",   "authNoPriv (no encryption)",    False),
            ("noAuthNoPriv", "noAuthNoPriv (no auth/enc)",    False),
        ]:
            phase = v3.get(mode_key, {})
            _add_bold_paragraph(doc, f"Execution Command ({label}):")
            _add_grey_box(doc, phase.get("command", "N/A"))
            doc.add_paragraph("")
            _add_bold_paragraph(doc, f"Executed Command Output ({label}):")
            _add_grey_box(doc, phase.get("output", "No output"))
            doc.add_paragraph("")

            success = phase.get("success", False)
            if phase.get("terminal_screenshot") and os.path.exists(phase["terminal_screenshot"]):
                if expect_success:
                    overview = (
                        "The DUT responded to authPriv SNMPv3 — secure communication confirmed."
                        if success else
                        "DUT did not respond to authPriv SNMPv3 — secure mode may be misconfigured."
                    )
                else:
                    overview = (
                        f"DUT responded to {mode_key} — insecure mode is allowed (FAIL)."
                        if success else
                        f"DUT rejected {mode_key} — insecure mode correctly disabled (PASS)."
                    )
                _add_screenshot_block(
                    doc,
                    f"TC10 — SNMPv3 {label} Terminal Output",
                    phase["terminal_screenshot"]
                )
                _add_bold_paragraph(doc, "Explanation:")
                doc.add_paragraph(overview)

            if phase.get("wireshark_screenshot") and os.path.exists(phase["wireshark_screenshot"]):
                ws_overview = (
                    "Wireshark confirms msgFlags show authentication and privacy bits set — authPriv enforced."
                    if (mode_key == "authPriv" and success) else
                    f"Wireshark packet analysis for {mode_key} mode shown above."
                )
                _add_screenshot_block(
                    doc,
                    f"TC10 — SNMPv3 {label} Wireshark Evidence",
                    phase["wireshark_screenshot"]
                )
                _add_bold_paragraph(doc, "Explanation:")
                doc.add_paragraph(ws_overview)
            _spacer(doc, 8)

        _add_bold_paragraph(doc, "d) Test Observations:")
        if tc10 == "PASS":
            doc.add_paragraph(
                "The DUT allows only SNMPv3 authPriv and rejects authNoPriv and noAuthNoPriv. "
                "Confidentiality and integrity of SNMP communication is ensured — marked PASS."
            )
        else:
            doc.add_paragraph(
                "The DUT permits weak SNMPv3 security levels (authNoPriv or noAuthNoPriv) "
                "or does not support authPriv correctly — marked FAIL."
            )
        _add_bold_paragraph(doc, "e) Evidence Provided:")
        doc.add_paragraph("Screenshots and Wireshark captures provided.")
        _spacer(doc)

    # ── Final result table (all 10 TCs) ──────────────────────────────────────

    def _result_table(self, doc, next_sec,
                      tc1, tc2, tc3, tc4,
                      tc5, tc6, tc7, tc8,
                      tc9, tc10):

        def remark(result, pass_msg, fail_msg, na_msg="Service not detected — test not applicable."):
            if result == "PASS": return pass_msg
            if result == "FAIL": return fail_msg
            return na_msg

        rows = [
            ("1",  "TC1_SSH_SECURE_CRYPTOGRAPHIC_ALGORITHMS",  tc1,
             remark(tc1,
                    "Only secure SSH algorithms are supported.",
                    "Weak or unsupported SSH algorithms detected.")),
            ("2",  "TC2_SSH_SECURE_COMMUNICATION",             tc2,
             remark(tc2,
                    "SSH communication established using approved cryptographic controls.",
                    "SSH communication uses weak or non-compliant parameters.")),
            ("3",  "TC3_SSH_WEAK_CIPHER_REJECTION",            tc3,
             remark(tc3,
                    "DUT rejects all weak SSH algorithm negotiation attempts.",
                    "One or more weak SSH algorithms were successfully negotiated.")),
            ("4",  "TC4_SSH_NO_ENCRYPTION_REJECTION",          tc4,
             remark(tc4,
                    "DUT rejects unencrypted SSH communication (none cipher).",
                    "Unencrypted SSH communication allowed — critical vulnerability.")),
            ("5",  "TC5_HTTPS_SECURE_CRYPTOGRAPHIC_ALGORITHMS", tc5,
             remark(tc5,
                    "Only secure TLS cipher suites are supported.",
                    "Weak or unsupported TLS cipher suites detected.")),
            ("6",  "TC6_HTTPS_SECURE_COMMUNICATION",           tc6,
             remark(tc6,
                    "HTTPS session established using approved TLS protocol and cipher.",
                    "HTTPS session uses weak or non-compliant TLS configuration.")),
            ("7",  "TC7_HTTPS_WEAK_CIPHER_REJECTION",          tc7,
             remark(tc7,
                    "DUT rejects all weak TLS cipher negotiation attempts.",
                    "One or more weak TLS ciphers were successfully negotiated.")),
            ("8",  "TC8_HTTPS_NO_ENCRYPTION_REJECTION",        tc8,
             remark(tc8,
                    "DUT enforces encryption and rejects NULL/unencrypted TLS communication.",
                    "Unencrypted TLS communication allowed — critical vulnerability.")),
            ("9",  "TC9_SNMP_VERSION_CHECK",                   tc9,
             remark(tc9,
                    "Only SNMPv3 is enabled; SNMPv1/v2c are disabled.",
                    "SNMPv1/v2c are enabled — insecure communication possible.")),
            ("10", "TC10_SNMP_SECURE_COMMUNICATION",           tc10,
             remark(tc10,
                    "SNMPv3 authPriv enforced; weaker modes (authNoPriv/noAuthNoPriv) rejected.",
                    "Weak SNMPv3 modes are allowed or authPriv is not correctly enforced.")),
        ]

        h = _add_itsar_heading(doc, f"{next_sec()}. Test Case Result Summary")
        _keep_with_next(h)

        rt = doc.add_table(rows=len(rows) + 1, cols=4)
        rt.style = "Table Grid"
        _style_header_row(rt, ["SL. No", "TEST CASE NAME", "PASS / FAIL", "Remarks"])

        for i, (sl, name, result, remark_text) in enumerate(rows, start=1):
            rt.cell(i, 0).text = sl
            rt.cell(i, 1).text = name
            rt.cell(i, 2).text = result
            rt.cell(i, 3).text = remark_text
            # Colour result cell
            res_cell = rt.cell(i, 2)
            color = "C6EFCE" if result == "PASS" else ("FCE4D6" if result == "FAIL" else "F2F2F2")
            tcPr = res_cell._tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:fill"), color)
            tcPr.append(shd)

        _pad_data_rows(rt)
        _prevent_row_split(rt)

    # ── Helpers ──────────────────────────────────────────────────────────────

    @staticmethod
    def _weak_result(weak_cipher_result: dict) -> str:
        for r in weak_cipher_result.get("results", []):
            if r.get("negotiated"):
                return "FAIL"
        return "PASS"

    @staticmethod
    def _https_weak_result(https_weak_result: dict) -> str:
        for r in https_weak_result.get("results", []):
            if r.get("negotiated"):
                return "FAIL"
        return "PASS"
